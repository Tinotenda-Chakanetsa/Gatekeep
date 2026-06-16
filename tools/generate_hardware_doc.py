#!/usr/bin/env python3
"""Generate the Gatekeeper hardware components & wiring guide as a .docx file.

Run with a Python that has python-docx installed:
    python tools/generate_hardware_doc.py
Output: Gatekeeper_Hardware_Guide.docx in the project root.

Covers the split two-board architecture: ESP32 control node (PIR + servo + LEDs +
buzzer) talking over UART to an ESP32-CAM camera node that handles the HTTPS upload
to the backend.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_PATH = os.path.join(PROJECT_ROOT, 'Gatekeeper_Hardware_Guide.docx')

ACCENT = RGBColor(0x2E, 0x6B, 0x4F)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(text)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def add_mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p


def build():
    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading('Gatekeeper', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('Hardware Components & Wiring Guide')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    tagline = doc.add_paragraph(
        'Automated licence-plate gate access control — split two-board prototype '
        '(ESP32 control node + ESP32-CAM camera node)'
    )
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].italic = True
    tagline.runs[0].font.color.rgb = ACCENT
    doc.add_paragraph()

    # ---- 1. Overview ----
    doc.add_heading('1. System overview', level=1)
    doc.add_paragraph(
        'In this build, two microcontrollers cooperate at the gate. The ESP32 dev board is '
        'the control node — it watches the PIR sensor and, when authorised, raises an SG90 '
        'servo boom while driving status LEDs and a buzzer. The ESP32-CAM is the camera node '
        '— it has the WiFi/HTTPS responsibility and the OV2640 camera. Splitting the work '
        'this way frees both boards from GPIO contention and removes the actuation lag that '
        'comes from running everything on the cam.'
    )
    doc.add_paragraph('Signal and command flow:')
    add_mono(
        doc,
        'Vehicle\n'
        '   |  (motion)\n'
        '   v\n'
        '[ PIR ] --> [ ESP32 control node ]\n'
        '                   |  "CAPTURE\\n" over UART\n'
        '                   v\n'
        '             [ ESP32-CAM ] ---- WiFi/HTTPS ----> [ Backend server ]\n'
        '                   ^                            YOLO + PaddleOCR\n'
        '                   |  "AUTH:1\\n" / "AUTH:0\\n" / "ERR\\n"\n'
        '             [ ESP32 control node ]\n'
        '                   |  drives:\n'
        '                   v\n'
        'Servo (boom)  +  Green/Red LEDs  +  Buzzer\n'
        '\n'
        'Dashboard (Vercel SPA) ---- HTTPS ----> Backend ----> "OPEN" pushed back to ESP32 via the cam',
    )

    # ---- 2. Bill of materials ----
    doc.add_heading('2. Bill of materials (what to buy)', level=1)
    doc.add_paragraph(
        'Total street cost: roughly USD 20 for the full split build. Items already in your '
        'kit are simply checked off. The single-channel relay listed at the bottom is OPTIONAL '
        'for this prototype because the SG90 servo simulates the boom directly — keep the '
        'relay aside for a future upgrade to a real mains-powered gate motor.'
    )
    add_table(
        doc,
        ['Component', 'Purpose', 'Qty', 'Approx. cost', 'Notes'],
        [
            ['ESP32 dev board (ESP32-WROOM-32, e.g. DevKit V1)', 'Control node: PIR + servo + LEDs + buzzer + UART link',
             '1', '$5–8', 'USB-C/Micro-USB powered; 30-pin generic layout matches the wiring tables below.'],
            ['AI-Thinker ESP32-CAM (OV2640)', 'Camera node: capture + WiFi/HTTPS to the backend',
             '1', '$7–10', 'No native USB — flashed via a USB-TTL adapter.'],
            ['USB-TTL programmer / ESP32-CAM-MB', 'One-time flashing of the ESP32-CAM',
             '1', '$2–4', 'The "ESP32-CAM-MB" dongle is the easiest option.'],
            ['PIR motion sensor (HC-SR501)', 'Trigger when a vehicle approaches',
             '1', '$1–3', 'Trim the on-board pots if false-triggering.'],
            ['SG90 micro servo', 'Drives the toll boom (simulated barrier)',
             '1', '$2–3', 'Must be powered from external 5 V, NOT the ESP32 3V3 pin.'],
            ['Green LED', 'GRANTED indicator', '1', '$0.10', 'Through a 220 Ω resistor to ground.'],
            ['Red LED', 'DENIED indicator', '1', '$0.10', 'Through a 220 Ω resistor to ground.'],
            ['220 Ω resistors', 'Current-limit the two LEDs', '2', '$0.20', 'One per LED, anode side.'],
            ['Active buzzer (3.3–5 V)', 'Audible grant/deny feedback', '1', '$1', 'A passive piezo also works with the tone() API.'],
            ['5 V 2 A power supply', 'Shared rail for both boards + servo + buzzer',
             '1', '$4–6', 'Use a clean 5 V supply; weak USB causes brownouts mid-capture.'],
            ['Jumper wires (M-F, F-F)', 'All inter-component wiring',
             '~15', '$2', 'Mix of M-F and F-F is convenient on a breadboard.'],
            ['Breadboard', 'Prototyping', '1', '$2–4', 'A half-size board is plenty.'],
            ['470 µF electrolytic capacitor (optional)', 'Bulk decoupling near the servo rail',
             '1', '$0.20', 'Strongly recommended — smooths the servo current spikes.'],
            ['1-channel 5 V relay module (OPTIONAL)', 'Only needed if you later drive a real gate motor instead of the SG90',
             '0–1', '$1–3', 'Leave disconnected for the prototype.'],
        ],
    )

    # ---- 3. Inter-board UART link ----
    doc.add_heading('3. Inter-board UART link (ESP32 ↔ ESP32-CAM)', level=1)
    doc.add_paragraph(
        'Three wires (plus a shared ground) carry every message between the two boards. The '
        'link runs at 115200 baud, 8N1, line-terminated text — exactly the same format used '
        'by the Arduino Serial Monitor. TX on one side connects to RX on the other (crossed).'
    )
    add_table(
        doc,
        ['ESP32-CAM pin', 'Direction', 'ESP32 pin', 'Notes'],
        [
            ['GPIO14 (labelled IO14)', '→ TX', 'D16 (RX2 / GPIO16)', 'Camera-node TX → control-node RX.'],
            ['GPIO15 (labelled IO15)', '← RX', 'D17 (TX2 / GPIO17)', 'Camera-node RX ← control-node TX.'],
            ['GND', '—', 'GND', 'Common ground is mandatory; without it UART will be unreliable.'],
            ['5V', '—', '5V / VIN', 'Both boards share the same 5 V rail.'],
        ],
    )
    doc.add_paragraph('Protocol (text lines, \\n-terminated):')
    add_mono(
        doc,
        'ESP32  -->  ESP32-CAM\n'
        '    CAPTURE\\n        ask the cam to take a picture and check it now\n'
        '\n'
        'ESP32-CAM  -->  ESP32\n'
        '    READY\\n          sent once at boot after WiFi connects\n'
        '    AUTH:1\\n         backend authorised this plate -> grantedFeedback()\n'
        '    AUTH:0\\n         backend denied (or no plate detected) -> deniedFeedback()\n'
        '    ERR\\n            capture or HTTP failed -> deniedFeedback()\n'
        '    OPEN\\n           asynchronous manual-open from the dashboard',
    )

    # ---- 4. ESP32 control-node pinout ----
    doc.add_heading('4. ESP32 control-node pinout (PIR, servo, LEDs, buzzer)', level=1)
    doc.add_paragraph(
        'These pin labels match a generic 30-pin ESP32 DevKit (and the labels visible on the '
        'board you flashed). All grounds tie to the shared GND rail.'
    )
    add_table(
        doc,
        ['Peripheral', 'ESP32 pin label', 'GPIO', 'Wiring detail'],
        [
            ['PIR OUT (signal)', 'D13', '13', 'Goes HIGH while motion is detected.'],
            ['Servo SG90 (signal — yellow/orange)', 'D25', '25', 'PWM-capable; servo red → 5 V rail, brown → GND.'],
            ['Green LED (+)', 'D26', '26', 'Anode → 220 Ω resistor → GPIO; cathode → GND.'],
            ['Red LED (+)', 'D27', '27', 'Anode → 220 Ω resistor → GPIO; cathode → GND.'],
            ['Buzzer (+)', 'D33', '33', 'Active buzzer; buzzer (–) to GND.'],
            ['UART link RX', 'RX2 (D16)', '16', 'Receives from ESP32-CAM GPIO14.'],
            ['UART link TX', 'TX2 (D17)', '17', 'Transmits to ESP32-CAM GPIO15.'],
            ['5 V supply', 'VIN / 5V', '—', 'From the shared 5 V rail (USB or external PSU).'],
            ['Ground', 'GND', '—', 'Tie to every other GND in the build.'],
        ],
    )
    caution = doc.add_paragraph()
    caution.add_run('Servo power caution: ').bold = True
    caution.add_run(
        'Never power the SG90 from the ESP32\'s 3V3 pin — the inrush current will brown out '
        'the regulator and crash the chip mid-WiFi. Feed it from the 5 V rail and add a 470 µF '
        'capacitor across V+/GND close to the servo if you can.'
    )

    # ---- 5. ESP32-CAM pinout ----
    doc.add_heading('5. ESP32-CAM camera-node pinout', level=1)
    doc.add_paragraph(
        'Only four pins are wired on the camera node — the rest are consumed by the OV2640 '
        'camera and the USB-TTL programmer footprint. No PIR, no relay, no LEDs on this board.'
    )
    add_table(
        doc,
        ['Function', 'ESP32-CAM pin', 'Notes'],
        [
            ['UART TX → ESP32 RX2', 'GPIO14 (IO14)', 'Used as HardwareSerial(1) TX in the sketch.'],
            ['UART RX ← ESP32 TX2', 'GPIO15 (IO15)', 'Used as HardwareSerial(1) RX in the sketch.'],
            ['Power', '5V', 'Shared 5 V rail with the ESP32 + servo.'],
            ['Ground', 'GND', 'Common with everything else.'],
        ],
    )

    # ---- 6. Flashing ----
    doc.add_heading('6. Flashing both boards', level=1)
    doc.add_paragraph(
        'The two firmware sketches live in separate folders. Flash each board with its own '
        'sketch; they have nothing to do with each other at build time.'
    )

    doc.add_heading('6.1 ESP32 control node — esp32_control_node.ino', level=2)
    for step in [
        'Arduino IDE → Tools → Manage Libraries → install "ESP32Servo" (Kevin Harrington / John K. Bennett).',
        'Board: "ESP32 Dev Module" (or whichever matches your specific dev board).',
        'Plug the ESP32 into USB. No IO0 jumper required — the dev board has a built-in USB-Serial bridge.',
        'Copy firmware/esp32_control_node/config.h.example → config.h (only timing parameters live here; no secrets).',
        'Click Upload. On Serial Monitor at 115200 baud you should see: "[ctrl] ready, listening for PIR + CAM".',
    ]:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('6.2 ESP32-CAM camera node — esp32cam_camera_node.ino', level=2)
    add_table(
        doc,
        ['USB-TTL / FTDI', 'ESP32-CAM', 'Notes'],
        [
            ['5V', '5V', 'Use 5V for reliable flashing.'],
            ['GND', 'GND', 'Common ground.'],
            ['TX', 'U0R (GPIO3 / RX)', 'Adapter transmit → board receive.'],
            ['RX', 'U0T (GPIO1 / TX)', 'Adapter receive → board transmit.'],
            ['— (jumper)', 'IO0  ↔  GND', 'Connect IO0 to GND to enter flash mode; remove afterwards.'],
        ],
    )
    for step in [
        'Arduino IDE → Board: "AI Thinker ESP32-CAM".',
        'Connect IO0 → GND and press the on-board RESET (or power-cycle) to enter flash mode.',
        'Copy firmware/esp32cam_camera_node/config.h.example → config.h. Fill in WiFi, '
        'SERVER_BASE_URL and DEVICE_API_KEY (from the dashboard Devices page).',
        'Click Upload. After "Done uploading", remove the IO0 → GND jumper and press RESET.',
        'Serial Monitor at 115200 baud should print "[wifi] connected: ..." then "[cam] READY".',
    ]:
        doc.add_paragraph(step, style='List Number')

    # ---- 7. Assembly ----
    doc.add_heading('7. Assembly checklist', level=1)
    for step in [
        'Flash both boards (Section 6) before connecting any peripherals.',
        'Power down. Wire the two boards together: 4 wires (TX, RX crossed, GND, 5V).',
        'Wire the PIR: OUT → ESP32 D13, VCC → 5 V, GND → GND.',
        'Wire the servo: signal → ESP32 D25, V+ → 5 V rail (with optional 470 µF cap), V– → GND.',
        'Wire the green LED: anode → 220 Ω → ESP32 D26, cathode → GND.',
        'Wire the red LED: anode → 220 Ω → ESP32 D27, cathode → GND.',
        'Wire the buzzer: (+) → ESP32 D33, (–) → GND.',
        'Confirm every module shares the same GND rail. This is the single most common cause of UART weirdness.',
        'Apply 5 V 2 A power. Open Serial Monitor on each board (115200 baud) and confirm READY/ctrl ready messages.',
        'Trigger the PIR (wave a hand) — the ESP32 should print "[ctrl] motion detected -> CAPTURE", '
        'the CAM should print "[link] CAPTURE received" then "[http] 200: {...}", and the boom should '
        'either lift (green LED + beep) or stay closed (red LED + 4 buzzes).',
    ]:
        doc.add_paragraph(step, style='List Number')

    # ---- 8. Power & safety ----
    doc.add_heading('8. Power & safety notes', level=1)
    for note in [
        'Use an external 5 V 2 A supply for the shared rail. The servo and the ESP32-CAM both '
        'spike current well above what a 500 mA USB port can deliver.',
        'Never power the SG90 servo from a 3V3 pin. Use 5 V plus a shared ground.',
        'Common ground is mandatory across every board, sensor and actuator. Floating grounds '
        'corrupt the UART link and cause spurious PIR triggers.',
        'The single-channel relay you have is optional for this prototype. Keep it for the '
        'day you replace the SG90 with a real gate motor — at which point the relay would '
        'switch the motor\'s contactor while the ESP32 stays at low voltage.',
        'The OV2640 camera is weak in low light. For night use add a white-LED or IR illuminator '
        'pointed at where the plate sits.',
    ]:
        doc.add_paragraph(note, style='List Bullet')

    # ---- 9. How the hardware ties into the software ----
    doc.add_heading('9. How the hardware connects to the software', level=1)
    doc.add_paragraph(
        'Only the ESP32-CAM speaks HTTP to the backend; the ESP32 control node never touches '
        'the network. The two endpoints involved are:'
    )
    add_table(
        doc,
        ['Action', 'Endpoint', 'Auth', 'Purpose'],
        [
            ['Send a capture', 'POST /api/gate/check/', 'X-Device-Key header',
             'Upload the plate photo; receive {"authorized": true|false}.'],
            ['Poll for commands', 'GET /api/gate/command/', 'X-Device-Key header',
             'Receive a manual "open" command from an operator; also acts as the online heartbeat.'],
        ],
    )
    steps = doc.add_paragraph()
    steps.add_run('Getting the device API key: ').bold = True
    steps.add_run(
        'log in to the dashboard as an administrator → Devices page → "+ Add Device" → click '
        '"Show" to reveal the key → paste it into DEVICE_API_KEY in the camera-node config.h.'
    )
    doc.add_paragraph(
        'No hardware yet? Use tools/gate_simulator.py to send a real image to the same '
        '/api/gate/check/ endpoint and demo the whole pipeline without any board.'
    )

    # ---- 10. Single-board alternative ----
    doc.add_heading('10. Single-board alternative (legacy)', level=1)
    doc.add_paragraph(
        'The original single-board layout (firmware/esp32cam_gate/esp32cam_gate.ino) is still '
        'in the repository. It runs everything on the ESP32-CAM with one relay output and one '
        'PIR input. It is functional but constrained by the cam\'s limited free GPIO and by '
        'the awkwardness of driving high-current peripherals from the cam\'s power rails. '
        'Prefer the split layout above for any build with a servo, LEDs and a buzzer.'
    )

    doc.save(OUTPUT_PATH)
    print(f'Wrote {OUTPUT_PATH}')


if __name__ == '__main__':
    build()
